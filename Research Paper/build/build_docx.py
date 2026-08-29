# -*- coding: utf-8 -*-
"""Render paper_content.py into an IEEE A4 two-column .docx.

Geometry is copied from templates/conference-template-a4.docx:
    page      595.30 x 841.90 pt  (A4)
    margins   top 54 pt, bottom 72 pt, left/right 45.35 pt
    body      2 columns, 18 pt gutter  ->  column width 243.3 pt
    text      Times New Roman 10 pt, justified, 0.2 in first-line indent
"""
import os
import re
import copy

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

import paper_content as C

#: Word records whoever saved the file; use this machine's account name.
MACHINE_USER = os.environ.get("USERNAME") or "user"

HERE = os.path.dirname(os.path.abspath(__file__))
PAPER_DIR = r"c:\Users\94775\Desktop\box\R26-IT-083\Research Paper"
FIG_DIR = os.path.join(PAPER_DIR, "figures")
OUT = os.path.join(PAPER_DIR, "Conference_Paper_DoubleBlind.docx")

PAGE_W, PAGE_H = 595.30, 841.90
M_TOP, M_BOT, M_SIDE = 54.0, 72.0, 45.35
GUTTER = 18.0
COL_W = (PAGE_W - 2 * M_SIDE - GUTTER) / 2.0          # 243.3 pt
FULL_W = PAGE_W - 2 * M_SIDE                          # 504.6 pt

ROMAN = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
ALPHA = "ABCDEFGHIJKLMNOP"


# ---------------------------------------------------------------- low level
def set_cols(section, num, space_pt=GUTTER):
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(round(space_pt * 20))))
    cols.set(qn("w:equalWidth"), "1")
    for c in cols.findall(qn("w:col")):
        cols.remove(c)


def set_geometry(section):
    section.page_width = Pt(PAGE_W)
    section.page_height = Pt(PAGE_H)
    section.top_margin = Pt(M_TOP)
    section.bottom_margin = Pt(M_BOT)
    section.left_margin = Pt(M_SIDE)
    section.right_margin = Pt(M_SIDE)
    section.header_distance = Pt(36)
    section.footer_distance = Pt(36)
    section.gutter = Pt(0)


def para(doc, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0,
         indent=0.0, line=None, keep=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Inches(indent)
    pf.left_indent = Inches(0)
    pf.right_indent = Inches(0)
    pf.widow_control = True
    if line is not None:
        pf.line_spacing = line
    if keep:
        pf.keep_with_next = True
    return p


def run(p, text, size=10, bold=False, italic=False, smallcaps=False,
        name="Times New Roman"):
    r = p.add_run(text)
    r.font.name = name
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.small_caps = smallcaps
    r.font.color.rgb = RGBColor(0, 0, 0)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)
    return r


TOKEN = re.compile(r"(\*[^*]+\*|\[\[\d+\]\]|\^\{[^}]*\})")


def rich(p, text, size=10, bold=False, base_italic=False):
    """Render inline markup: *italic* and [[n]] citations."""
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            run(p, piece[1:-1], size=size, bold=bold, italic=True)
        elif piece.startswith("[[") and piece.endswith("]]"):
            run(p, "[%s]" % piece[2:-2], size=size, bold=bold,
                italic=base_italic)
        elif piece.startswith("^{") and piece.endswith("}"):
            r = run(p, piece[2:-1], size=size, bold=bold, italic=base_italic)
            r.font.superscript = True
        else:
            run(p, piece, size=size, bold=bold, italic=base_italic)


def cell_borders(cell, top=None, bottom=None, grid=False):
    """grid=True draws all four edges, as in the reference paper's tables."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    if grid:
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement("w:%s" % edge)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tcPr.append(borders)
        return
    for edge, width in (("top", top), ("bottom", bottom)):
        el = OxmlElement("w:%s" % edge)
        if width:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(width))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    tcPr.append(el)


def cell_margins(table, top=1.2, bottom=1.2, left=2.5, right=2.5):
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for name, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement("w:%s" % name)
        el.set(qn("w:w"), str(int(val * 20)))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)



def fixed_layout(table, total_pt):
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(total_pt * 20)))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)



# ------------------------------------------------- markup sanity guard
def check_markup():
    """A stray '*' pairs with the next one and swallows whole sentences."""
    bad = []
    for kind, payload in C.BODY:
        texts = []
        if kind == "P":
            texts = [payload]
        elif kind == "LIST":
            texts = payload
        elif kind == "FIG":
            texts = [payload[1]]
        elif kind == "TABLE":
            texts = [payload["caption"]] + [c for r in payload["rows"]
                                            for c in r]
        for t in texts:
            for piece in TOKEN.split(t):
                if piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
                    body = piece[1:-1]
                    if len(body) > 40 or "[[" in body:
                        bad.append(body[:70])
                elif piece and not piece.startswith("[[")                         and not piece.startswith("^{") and "[[" in piece:
                    bad.append(piece[:70])
    if bad:
        raise SystemExit("markup would render wrongly: " + "; ".join(bad))


check_markup()

# ------------------------------------------------------------------ builder
doc = Document()

# base style
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10)
rpr = normal.element.get_or_add_rPr()
rf = rpr.find(qn("w:rFonts"))
if rf is None:
    rf = OxmlElement("w:rFonts")
    rpr.insert(0, rf)
for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rf.set(qn(a), "Times New Roman")
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

sec0 = doc.sections[0]
set_geometry(sec0)
set_cols(sec0, 1)

# ----------------------------------------------------------------- title
p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
run(p, C.TITLE, size=24)

# ------------------------------------------------------------ author block
if getattr(C, "ANONYMOUS", False):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=0)
    run(p, "Anonymous Author(s)", size=11)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    run(p, "Affiliation withheld for double-blind review", size=10,
        italic=True)
else:
    atab = doc.add_table(rows=1, cols=len(C.AUTHORS))
    atab.alignment = WD_TABLE_ALIGNMENT.CENTER
    atab.autofit = False
    cell_margins(atab, 0, 0, 0, 0)
    fixed_layout(atab, FULL_W)
    for i, (name, dept, org, city, sid) in enumerate(C.AUTHORS):
        cell = atab.rows[0].cells[i]
        cell.width = Pt(FULL_W / len(C.AUTHORS))
        cell_borders(cell)
        first = True
        for text, sz, it in ((name, 11, False), (dept, 10, True),
                             (org, 10, True), (city, 10, True),
                             (sid, 10, False)):
            cp = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.line_spacing = 1.0
            run(cp, text, size=sz, italic=it)

# ------------------------------------------------- two-column body section
body = doc.add_section(WD_SECTION.CONTINUOUS)
set_geometry(body)
set_cols(body, 2)
body.start_type = WD_SECTION.CONTINUOUS

# abstract
p = para(doc, before=10, after=0, indent=0.19)
run(p, "Abstract\u2014", size=9, bold=True, italic=True)
rich(p, C.ABSTRACT, size=9, bold=True)

p = para(doc, before=6, after=8, indent=0.19)
run(p, "Index Terms\u2014", size=9, bold=True, italic=True)
run(p, C.INDEX_TERMS, size=9, bold=True)

state = {"h1": 0, "h2": 0, "eq": 0, "fig": 0, "tab": 0, "in_span": False}


def open_span():
    """Switch to a single-column continuous section for a wide float."""
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    set_geometry(s)
    set_cols(s, 1)
    state["in_span"] = True


def close_span():
    s = doc.add_section(WD_SECTION.CONTINUOUS)
    set_geometry(s)
    set_cols(s, 2)
    state["in_span"] = False


def add_figure(fname, caption, span):
    if span:
        open_span()
    width = Inches(FULL_W / 72.0) if span else Inches(COL_W / 72.0)
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
    p.add_run().add_picture(os.path.join(FIG_DIR, fname), width=width)
    state["fig"] += 1
    cap = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)
    run(cap, "Fig. %d. " % state["fig"], size=8)
    rich(cap, caption, size=8)
    if span:
        close_span()


def add_table(spec):
    span = spec.get("span", False)
    if span:
        open_span()
    state["tab"] += 1
    cap = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=2,
               keep=True)
    run(cap, "TABLE %s" % ROMAN[state["tab"]], size=8)
    cap2 = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=3, keep=True)
    run(cap2, spec["caption"], size=8, smallcaps=True)

    total = FULL_W if span else COL_W
    widths = spec["widths"]
    t = doc.add_table(rows=1, cols=len(spec["cols"]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell_margins(t)
    fixed_layout(t, total)
    hdr = t.rows[0]
    for j, name in enumerate(spec["cols"]):
        c = hdr.cells[j]
        c.width = Pt(total * widths[j])
        cell_borders(c, grid=True)
        cp = c.paragraphs[0]
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.0
        run(cp, name, size=8, smallcaps=True)
    for i, row in enumerate(spec["rows"]):
        last = (i == len(spec["rows"]) - 1)
        tr = t.add_row()
        for j, val in enumerate(row):
            c = tr.cells[j]
            c.width = Pt(total * widths[j])
            cell_borders(c, grid=True)
            cp = c.paragraphs[0]
            cp.paragraph_format.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if j < 3 and len(val) < 12
                else WD_ALIGN_PARAGRAPH.LEFT)
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.line_spacing = 1.0
            rich(cp, val, size=8)
    if "notes" in spec:
        n = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=2, after=8)
        rich(n, spec["notes"], size=8)
    else:
        para(doc, after=0, before=4)
    if span:
        close_span()


for kind, payload in C.BODY:
    if kind == "H1":
        state["h1"] += 1
        state["h2"] = 0
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4,
                 keep=True)
        run(p, "%s. %s" % (ROMAN[state["h1"]], payload), size=10,
            smallcaps=True)
    elif kind == "H2":
        state["h2"] += 1
        p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=3,
                 keep=True)
        run(p, "%s. %s" % (ALPHA[state["h2"] - 1], payload), size=10,
            italic=True)
    elif kind == "P":
        p = para(doc, indent=0.2)
        rich(p, payload)
    elif kind == "LIST":
        for item in payload:
            p = para(doc, indent=0.0)
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            run(p, "\u2022  ")
            rich(p, item)
    elif kind == "EQ":
        state["eq"] += 1
        uni, _ = payload
        p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=4)
        w = COL_W
        p.paragraph_format.tab_stops.add_tab_stop(
            Pt(w / 2.0), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(
            Pt(w), WD_TAB_ALIGNMENT.RIGHT)
        run(p, "\t")
        run(p, uni, size=10, italic=True)
        run(p, "\t")
        run(p, "(%d)" % state["eq"], size=10)
    elif kind == "FIG":
        add_figure(*payload)
    elif kind == "TABLE":
        add_table(payload)

# ------------------------------------------------------------- references
p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, keep=True)
run(p, "References", size=10, smallcaps=True)
for i, ref in enumerate(C.REFERENCES, 1):
    p = para(doc, after=1)
    p.paragraph_format.left_indent = Inches(0.24)
    p.paragraph_format.first_line_indent = Inches(-0.24)
    run(p, "[%d]\t" % i, size=8)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(0.24))
    run(p, ref, size=8)

cp = doc.core_properties
cp.author = MACHINE_USER
cp.last_modified_by = MACHINE_USER
cp.title = ""
cp.subject = ""
cp.comments = ""
cp.category = ""
cp.keywords = ""
cp.identifier = ""

doc.save(OUT)
print("saved", OUT)

# ------------------------------------------------------ crude length check
chars = 0
for kind, payload in C.BODY:
    if kind == "P":
        chars += len(payload)
    elif kind == "LIST":
        chars += sum(len(x) for x in payload)
print("body prose characters:", chars,
      "-> approx", round(chars / 52.0), "lines of text")
